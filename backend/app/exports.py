"""
Génération des exports Word (.docx) et Excel (.xlsx) de Djeliya.

Deux niveaux :
- Entretien seul : transcription + son analyse.
- Étude (corpus entier) : rapport de recherche complet — page de garde,
  sommaire, démarche méthodologique référencée (APA), structure des
  résultats numérotée, fiabilité inter-codeurs, annexes (liste des
  entretiens et transcriptions intégrales).

Bilingue (fr/en) : la langue du rapport suit la langue dans laquelle
l'analyse IA a été générée (entretien/corpus "analyse_langue"), pour que
les libellés fixes du document restent cohérents avec le contenu généré.
Les verbatims eux-mêmes ne sont jamais traduits (voir main.py).
"""

import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
from .charts_quant import graphique_fiabilite, graphique_moyennes, graphique_matrice_correlations

OR_HEX = "E4B04A"

REFERENCES_APA = {
    "gioia": (
        "Gioia, D. A., Corley, K. G., & Hamilton, A. L. (2013). Seeking qualitative "
        "rigor in inductive research: Notes on the Gioia methodology. Organizational "
        "Research Methods, 16(1), 15-31."
    ),
    "thematique": (
        "Braun, V., & Clarke, V. (2006). Using thematic analysis in psychology. "
        "Qualitative Research in Psychology, 3(2), 77-101.\n"
        "Braun, V., & Clarke, V. (2019). Reflecting on reflexive thematic analysis. "
        "Qualitative Research in Sport, Exercise and Health, 11(4), 589-597."
    ),
    "contenu": "Bardin, L. (2013). L'analyse de contenu (2e éd.). Paris : Presses Universitaires de France.",
}

METHODE_LABEL = {
    "fr": {
        "gioia": "Méthode de structuration des données de Gioia et al. (2013)",
        "thematique": "Analyse thématique réflexive de Braun & Clarke (2006, 2019)",
        "contenu": "Analyse de contenu catégorielle de Bardin (2013)",
    },
    "en": {
        "gioia": "Gioia et al. (2013) data structuring method",
        "thematique": "Braun & Clarke (2006, 2019) reflexive thematic analysis",
        "contenu": "Bardin (2013) categorical content analysis",
    },
}

# ----------------------------------------------------------------- libellés bilingues
L = {
    "fr": {
        "sous_titre_entretien": "Transcription et analyse qualitative",
        "langue_detectee": "Langue détectée : ",
        "genere_le": "Document généré par Djeliya le ",
        "sommaire_repli": "Sommaire — clic droit puis « Mettre à jour les champs » pour l'afficher.",
        "page": "Page ", "sur": " / ",
        "transcription": "1. Transcription",
        "analyse_qualitative": "2. Analyse qualitative",
        "methode": "Méthode : ", "modele": "Modèle : ",
        "demarche_methodologique": "Démarche méthodologique",
        "structure_resultats": "Structure des résultats",
        "synthese_interpretative": "Synthèse interprétative",
        "limites_analyse": "Limites de l'analyse automatique",
        "reference": "Référence", "references": "Références",
        "rapport_titre": "RAPPORT D'ANALYSE QUALITATIVE",
        "etude": "Étude",
        "entretiens_mot": "entretien(s)", "materiau_audio": " de matériau audio",
        "langues_parentheses": "Langue(s) : ",
        "presentation_etude": "1. Présentation de l'étude",
        "etude_intro": "Cette étude qualitative repose sur un corpus de {n} entretien(s), totalisant {duree} de matériau audio transcrit.",
        "question_recherche": "Question de recherche / angle d'analyse : ",
        "entretien_col": "Entretien", "langue_col": "Langue", "duree_col": "Durée", "date_col": "Date",
        "demarche_titre2": "2. Démarche méthodologique",
        "aucune_analyse_corpus": "Aucune analyse transversale n'a encore été lancée sur ce corpus.",
        "structure_resultats3": "3. Structure des résultats",
        "synthese4": "4. Synthèse interprétative",
        "limites5": "5. Limites méthodologiques",
        "fiabilite6": "6. Fiabilité inter-codeurs",
        "fiabilite_intro": "Le kappa de Cohen mesure l'accord entre codeurs indépendants au-delà de ce que le hasard produirait seul (0 = accord aléatoire, 1 = accord parfait).",
        "codeurs_col": "Codeurs", "kappa_moyen_col": "Kappa moyen", "interpretation_col": "Interprétation",
        "rigueur7": "7. Indicateurs de rigueur qualitative",
        "convergence71": "7.1. Convergence entre entretiens",
        "convergence_intro": "Un thème retrouvé dans plusieurs entretiens distincts est plus robuste qu'un thème isolé à un seul cas — c'est le principe de triangulation en recherche qualitative.",
        "theme_col": "Thème", "entretiens_concernes_col": "Entretiens concernés", "pct_corpus_col": "% du corpus",
        "saturation72": "7.2. Saturation théorique",
        "saturation_intro": "Nombre de concepts nouveaux apportés par chaque entretien, dans l'ordre chronologique. Une courbe qui s'aplatit vers la fin du corpus suggère que des entretiens supplémentaires n'apporteraient plus beaucoup d'éléments inédits.",
        "concepts_nouveaux_col": "Concepts nouveaux", "cumul_col": "Cumul",
        "annexe_a": "Annexe A — Composition du corpus",
        "annexe_b": "Annexe B — Transcriptions intégrales",
        "titre_col": "Titre",
        "locuteur_prefixe": "Locuteur ",
        # Excel
        "feuille_transcription": "Transcription", "feuille_statistiques": "Statistiques",
        "feuille_analyse": "Analyse", "feuille_vue_ensemble": "Vue d'ensemble",
        "feuille_entretiens": "Entretiens", "feuille_transcriptions": "Transcriptions",
        "feuille_analyse_transversale": "Analyse transversale", "feuille_codebook": "Codebook",
        "feuille_stats_qual": "Statistiques qualitatives", "feuille_fiabilite": "Fiabilité inter-codeurs",
        "debut": "Début", "fin": "Fin", "locuteur": "Locuteur", "texte": "Texte",
        "confiance_moy": "Confiance moyenne (%)", "confiance_pct": "Fiabilité (%)",
        "indicateur": "Indicateur", "valeur": "Valeur",
        "titre_stat": "Titre", "langue_detectee_stat": "Langue détectée", "duree_totale": "Durée totale",
        "nb_segments": "Nombre de segments", "nb_mots": "Nombre de mots transcrits",
        "fiabilite_moy": "Fiabilité moyenne (%)", "temps_parole": "Temps de parole par locuteur",
        "niveau": "Niveau", "libelle": "Libellé", "detail_verbatim": "Détail / verbatim",
        "horodatage": "Horodatage", "premier_ordre": "Premier ordre", "second_ordre": "Second ordre",
        "dimension_agregee": "Dimension agrégée", "dimension": "Dimension", "theme": "Thème",
        "concept": "Concept", "frequence": "Fréquence (verbatims)",
        "freq_dimension": "Fréquence par dimension", "themes_pl": "Thèmes", "concepts_pl": "Concepts",
        "verbatims_pl": "Verbatims", "convergence_titre": "Convergence entre entretiens (triangulation)",
        "saturation_titre": "Saturation théorique (ordre chronologique)",
        "concepts_distincts": "Cumul de concepts distincts",
        "etude_corpus": "Étude / corpus", "methodologie": "Méthodologie",
        "question_recherche_stat": "Question de recherche", "nb_entretiens": "Nombre d'entretiens",
        "duree_cumulee": "Durée cumulée", "nb_mots_total": "Nombre total de mots transcrits",
        "fiabilite_moy_corpus": "Fiabilité moyenne du corpus (%)", "rapport_genere_le": "Rapport généré le",
        "segments": "Segments", "mots": "Mots", "date": "Date",
        # Guide d'entretien
        "guide_sous_titre": "Guide d'entretien de recherche",
        "guide_theme": "Thème : ", "guide_question": "Question de recherche : ",
        "guide_infos_pratiques": "Informations pratiques",
        "guide_type": "Type d'entretien", "guide_duree": "Durée estimée",
        "guide_population": "Population cible", "guide_materiel": "Matériel recommandé",
        "guide_preambule": "Préambule à lire au participant",
        "guide_objectif": "Objectif : ",
        "guide_conseils": "Conseils méthodologiques pour la conduite de l'entretien",
        "guide_note": "Note méthodologique",
        "guide_relances": "Relances possibles :",
        # Étude quantitative
        "eq_sous_titre": "Cadre théorique, revue de littérature et méthodologie quantitative",
        "eq_theme": "Thème : ", "eq_question": "Question de recherche : ",
        "eq_cadre": "1. Cadre théorique", "eq_revue": "2. Revue de littérature",
        "eq_methodo": "3. Méthodologie", "eq_type_etude": "Type d'étude",
        "eq_population": "Population cible", "eq_echantillon": "Échantillon",
        "eq_hypotheses": "Hypothèses de recherche", "eq_hyp_col": "Code", "eq_hyp_enonce": "Énoncé",
        "eq_variables": "Variables du modèle", "eq_var_nom": "Variable", "eq_var_type": "Type", "eq_var_def": "Définition",
        "eq_questionnaire": "4. Questionnaire", "eq_note": "Note méthodologique",
        "eq_references": "5. Références", "eq_ref_methodo_intro": "Références méthodologiques vérifiées :",
        "eq_ref_concepts_intro": "Concepts théoriques mobilisés — références précises à compléter par le chercheur :",
        "eq_item_type": "Type", "eq_item_options": "Modalités / échelle",
        # Gabarit Excel
        "eq_feuille_reponses": "Réponses", "eq_feuille_guide": "Guide de saisie",
        "eq_guide_code": "Code", "eq_guide_libelle": "Libellé", "eq_guide_type": "Type de réponse attendue",
        "eq_guide_valeurs": "Valeurs possibles", "eq_guide_intro": "Une ligne = un·e répondant·e. Remplis la feuille « Réponses » en respectant les codes de colonnes ci-dessous ; consulte cette feuille pour connaître le type de réponse attendu pour chaque question.",
        "eq_type_libelle": {
            "choix_unique": "Choix unique parmi les modalités listées",
            "choix_multiple": "Un ou plusieurs choix, séparés par une virgule",
            "echelle_likert": "Nombre entier sur l'échelle indiquée",
            "numerique": "Valeur numérique",
            "texte_libre": "Texte libre",
        },
        # Analyse quantitative
        "aq_sous_titre": "Résultats de l'analyse statistique",
        "aq_apercu": "Aperçu de l'échantillon", "aq_n_repondants": "Nombre de répondants",
        "aq_descriptives": "1. Statistiques descriptives", "aq_desc_col": "Item", "aq_desc_n": "n",
        "aq_desc_moy": "Moyenne", "aq_desc_et": "Écart-type", "aq_desc_min": "Min", "aq_desc_max": "Max",
        "aq_desc_mediane": "Médiane", "aq_desc_ic95": "IC 95% de la moyenne",
        "aq_desc_asymetrie": "Asymétrie", "aq_desc_aplatissement": "Aplatissement",
        "aq_frequences": "2. Tableaux de fréquences", "aq_freq_modalite": "Modalité", "aq_freq_effectif": "Effectif", "aq_freq_pct": "%",
        "aq_fiabilite": "3. Fiabilité des construits (alpha de Cronbach)", "aq_fiab_intro": "L'alpha de Cronbach mesure la cohérence interne d'un ensemble d'items censés mesurer le même construit (0 = aucune cohérence, 1 = cohérence parfaite). Un seuil ≥ 0,70 est généralement considéré satisfaisant en sciences sociales.",
        "aq_fiab_variable": "Construit", "aq_fiab_nb_items": "Nb items", "aq_fiab_alpha": "Alpha", "aq_fiab_interpretation": "Interprétation",
        "aq_fiab_detail_titre": "Détail par item", "aq_fiab_item_total": "Corrélation item-total corrigée", "aq_fiab_alpha_supprime": "Alpha si item supprimé",
        "aq_fiab_normalite": "Normalité du score composite (Shapiro-Wilk)", "aq_fiab_normalite_oui": "distribution normale (p ≥ 0,05)",
        "aq_fiab_normalite_non": "distribution non normale (p < 0,05)",
        "aq_correlations": "4. Corrélations entre construits", "aq_corr_intro": "Corrélation de Pearson entre les scores composites de chaque construit, avec test de significativité.",
        "aq_afe_titre": "5. Analyse factorielle exploratoire (AFE)",
        "aq_afe_kmo": "Indice KMO (adéquation de l'échantillonnage)", "aq_afe_bartlett": "Test de sphéricité de Bartlett",
        "aq_afe_facteurs": "Facteurs extraits (critère de Kaiser)", "aq_afe_variance": "Variance expliquée cumulée",
        "aq_afe_charges": "Charges factorielles (après rotation varimax)",
        "aq_afc_titre": "6. Analyse factorielle confirmatoire (AFC)",
        "aq_afc_intro": "Teste si la structure de mesure postulée (chaque construit mesuré par ses items déclarés) s'ajuste bien aux données observées.",
        "aq_afc_indisponible": "L'AFC n'a pas pu être calculée pour cette analyse (modèle non convergent ou données insuffisantes) — les autres résultats restent valides.",
        "aq_reg_titre": "7. Régressions multiples",
        "aq_reg_intro": "Effet de chaque prédicteur sur la variable dépendante, en contrôlant simultanément les autres prédicteurs (coefficients standardisés, directement comparables entre eux).",
        "aq_reg_predicteur": "Prédicteur", "aq_reg_beta": "β standardisé", "aq_reg_p": "p", "aq_reg_ic": "IC95%",
        "aq_med_titre": "8. Tests de médiation (bootstrap)",
        "aq_med_intro": "Effet indirect (a × b) ré-estimé sur 2000 rééchantillonnages — la médiation est considérée significative lorsque l'intervalle de confiance à 95% exclut zéro (méthode de Preacher & Hayes).",
        "aq_med_chemin": "Chemin testé", "aq_med_indirect": "Effet indirect", "aq_med_ic": "IC95%", "aq_med_verdict": "Verdict",
        "aq_corr_v1": "Variable 1", "aq_corr_v2": "Variable 2", "aq_corr_r": "r", "aq_corr_p": "p", "aq_corr_methode": "Méthode",
        "aq_feuille_desc": "Statistiques descriptives", "aq_feuille_freq": "Tableaux de fréquences",
        "aq_feuille_fiab": "Fiabilité (alpha)", "aq_feuille_corr": "Corrélations",
        "aq_graphiques": "Graphiques", "aq_graph_fiabilite": "Fiabilité des construits",
        "aq_graph_moyennes": "Scores moyens par construit", "aq_graph_matrice": "Matrice de corrélations",
        "aq_synthese_titre": "9. Synthèse et interprétation", "aq_synthese_generale": "Synthèse générale",
        "aq_synthese_fiabilite": "Discussion de la fiabilité", "aq_synthese_hypotheses": "Test des hypothèses de recherche",
        "aq_synthese_limites": "Limites de l'analyse", "aq_synthese_recommandations": "Recommandations",
        "aq_hyp_code": "Hypothèse", "aq_hyp_verdict": "Verdict", "aq_hyp_justif": "Justification",
        "aq_synthese_avertissement": "Interprétation générée par IA à partir des résultats statistiques ci-dessus — à valider par le chercheur. Les corrélations rapportées n'établissent jamais formellement une médiation ou une modération, seulement une cohérence ou non avec l'hypothèse testée.",
        "aq_synthese_indisponible": "La synthèse interprétative n'a pas pu être générée pour cette analyse — les résultats statistiques bruts ci-dessus restent entièrement valides et exploitables.",
        "guide_grille_titre": "Grille de cohérence",
        "guide_grille_intro": "Chaque question principale du guide, mise en correspondance avec la dimension théorique qu'elle vise à explorer — à vérifier et amender par le chercheur.",
        "guide_grille_question": "Question", "guide_grille_dimension": "Dimension visée", "guide_grille_justif": "Justification",
        "avertissement_ia": (
            "Ce document a été généré automatiquement par intelligence artificielle. C'est un outil "
            "d'aide à la préparation — il doit être relu, validé et adapté par le chercheur avant toute "
            "utilisation sur le terrain ou dans un travail académique. Il ne constitue pas un instrument "
            "scientifiquement validé au sens psychométrique du terme."
        ),
    },
    "en": {
        "sous_titre_entretien": "Transcript and qualitative analysis",
        "langue_detectee": "Detected language: ",
        "genere_le": "Document generated by Djeliya on ",
        "sommaire_repli": "Table of contents — right-click then “Update Field” to display it.",
        "page": "Page ", "sur": " of ",
        "transcription": "1. Transcript",
        "analyse_qualitative": "2. Qualitative analysis",
        "methode": "Method: ", "modele": "Model: ",
        "demarche_methodologique": "Methodological approach",
        "structure_resultats": "Findings structure",
        "synthese_interpretative": "Interpretive summary",
        "limites_analyse": "Limits of the automated analysis",
        "reference": "Reference", "references": "References",
        "rapport_titre": "QUALITATIVE ANALYSIS REPORT",
        "etude": "Study",
        "entretiens_mot": "interview(s)", "materiau_audio": " of audio material",
        "langues_parentheses": "Language(s): ",
        "presentation_etude": "1. Study overview",
        "etude_intro": "This qualitative study is based on a corpus of {n} interview(s), totalling {duree} of transcribed audio material.",
        "question_recherche": "Research question / analysis angle: ",
        "entretien_col": "Interview", "langue_col": "Language", "duree_col": "Duration", "date_col": "Date",
        "demarche_titre2": "2. Methodological approach",
        "aucune_analyse_corpus": "No cross-case analysis has been run on this corpus yet.",
        "structure_resultats3": "3. Findings structure",
        "synthese4": "4. Interpretive summary",
        "limites5": "5. Methodological limitations",
        "fiabilite6": "6. Inter-rater reliability",
        "fiabilite_intro": "Cohen's kappa measures agreement between independent coders beyond what chance alone would produce (0 = random agreement, 1 = perfect agreement).",
        "codeurs_col": "Coders", "kappa_moyen_col": "Average kappa", "interpretation_col": "Interpretation",
        "rigueur7": "7. Qualitative rigor indicators",
        "convergence71": "7.1. Convergence across interviews",
        "convergence_intro": "A theme found across several distinct interviews is more robust than one isolated to a single case — this is the triangulation principle in qualitative research.",
        "theme_col": "Theme", "entretiens_concernes_col": "Interviews involved", "pct_corpus_col": "% of corpus",
        "saturation72": "7.2. Theoretical saturation",
        "saturation_intro": "Number of new concepts introduced by each interview, in chronological order. A curve flattening toward the end of the corpus suggests additional interviews would add little new material.",
        "concepts_nouveaux_col": "New concepts", "cumul_col": "Cumulative",
        "annexe_a": "Appendix A — Corpus composition",
        "annexe_b": "Appendix B — Full transcripts",
        "titre_col": "Title",
        "locuteur_prefixe": "Speaker ",
        # Excel
        "feuille_transcription": "Transcript", "feuille_statistiques": "Statistics",
        "feuille_analyse": "Analysis", "feuille_vue_ensemble": "Overview",
        "feuille_entretiens": "Interviews", "feuille_transcriptions": "Transcripts",
        "feuille_analyse_transversale": "Cross-case analysis", "feuille_codebook": "Codebook",
        "feuille_stats_qual": "Qualitative statistics", "feuille_fiabilite": "Inter-rater reliability",
        "debut": "Start", "fin": "End", "locuteur": "Speaker", "texte": "Text",
        "confiance_moy": "Average confidence (%)", "confiance_pct": "Reliability (%)",
        "indicateur": "Indicator", "valeur": "Value",
        "titre_stat": "Title", "langue_detectee_stat": "Detected language", "duree_totale": "Total duration",
        "nb_segments": "Number of segments", "nb_mots": "Number of transcribed words",
        "fiabilite_moy": "Average reliability (%)", "temps_parole": "Speaking time per speaker",
        "niveau": "Level", "libelle": "Label", "detail_verbatim": "Detail / verbatim",
        "horodatage": "Timestamp", "premier_ordre": "First order", "second_ordre": "Second order",
        "dimension_agregee": "Aggregate dimension", "dimension": "Dimension", "theme": "Theme",
        "concept": "Concept", "frequence": "Frequency (verbatims)",
        "freq_dimension": "Frequency by dimension", "themes_pl": "Themes", "concepts_pl": "Concepts",
        "verbatims_pl": "Verbatims", "convergence_titre": "Convergence across interviews (triangulation)",
        "saturation_titre": "Theoretical saturation (chronological order)",
        "concepts_distincts": "Cumulative distinct concepts",
        "etude_corpus": "Study / corpus", "methodologie": "Methodology",
        "question_recherche_stat": "Research question", "nb_entretiens": "Number of interviews",
        "duree_cumulee": "Cumulative duration", "nb_mots_total": "Total transcribed words",
        "fiabilite_moy_corpus": "Average corpus reliability (%)", "rapport_genere_le": "Report generated on",
        "segments": "Segments", "mots": "Words", "date": "Date",
        # Guide d'entretien
        "guide_sous_titre": "Research interview guide",
        "guide_theme": "Theme: ", "guide_question": "Research question: ",
        "guide_infos_pratiques": "Practical information",
        "guide_type": "Interview type", "guide_duree": "Estimated duration",
        "guide_population": "Target population", "guide_materiel": "Recommended materials",
        "guide_preambule": "Preamble to read to the participant",
        "guide_objectif": "Objective: ",
        "guide_conseils": "Methodological guidance for conducting the interview",
        "guide_note": "Methodological note",
        "guide_relances": "Possible follow-ups:",
        # Étude quantitative
        "eq_sous_titre": "Theoretical framework, literature review and quantitative methodology",
        "eq_theme": "Theme: ", "eq_question": "Research question: ",
        "eq_cadre": "1. Theoretical framework", "eq_revue": "2. Literature review",
        "eq_methodo": "3. Methodology", "eq_type_etude": "Study type",
        "eq_population": "Target population", "eq_echantillon": "Sample",
        "eq_hypotheses": "Research hypotheses", "eq_hyp_col": "Code", "eq_hyp_enonce": "Statement",
        "eq_variables": "Model variables", "eq_var_nom": "Variable", "eq_var_type": "Type", "eq_var_def": "Definition",
        "eq_questionnaire": "4. Questionnaire", "eq_note": "Methodological note",
        "eq_references": "5. References", "eq_ref_methodo_intro": "Verified methodological references:",
        "eq_ref_concepts_intro": "Theoretical concepts used — precise references to be completed by the researcher:",
        "eq_item_type": "Type", "eq_item_options": "Response options / scale",
        # Gabarit Excel
        "eq_feuille_reponses": "Responses", "eq_feuille_guide": "Data entry guide",
        "eq_guide_code": "Code", "eq_guide_libelle": "Label", "eq_guide_type": "Expected answer type",
        "eq_guide_valeurs": "Possible values", "eq_guide_intro": "One row = one respondent. Fill in the “Responses” sheet following the column codes below; refer to this sheet to know the expected answer type for each question.",
        "eq_type_libelle": {
            "choix_unique": "Single choice among the listed options",
            "choix_multiple": "One or more choices, comma-separated",
            "echelle_likert": "Whole number on the indicated scale",
            "numerique": "Numeric value",
            "texte_libre": "Free text",
        },
        # Analyse quantitative
        "aq_sous_titre": "Statistical analysis results",
        "aq_apercu": "Sample overview", "aq_n_repondants": "Number of respondents",
        "aq_descriptives": "1. Descriptive statistics", "aq_desc_col": "Item", "aq_desc_n": "n",
        "aq_desc_moy": "Mean", "aq_desc_et": "Std. dev.", "aq_desc_min": "Min", "aq_desc_max": "Max",
        "aq_desc_mediane": "Median", "aq_desc_ic95": "95% CI of the mean",
        "aq_desc_asymetrie": "Skewness", "aq_desc_aplatissement": "Kurtosis",
        "aq_frequences": "2. Frequency tables", "aq_freq_modalite": "Category", "aq_freq_effectif": "Count", "aq_freq_pct": "%",
        "aq_fiabilite": "3. Construct reliability (Cronbach's alpha)", "aq_fiab_intro": "Cronbach's alpha measures the internal consistency of a set of items intended to measure the same construct (0 = no consistency, 1 = perfect consistency). A threshold ≥ 0.70 is generally considered satisfactory in social sciences.",
        "aq_fiab_variable": "Construct", "aq_fiab_nb_items": "Nb. items", "aq_fiab_alpha": "Alpha", "aq_fiab_interpretation": "Interpretation",
        "aq_fiab_detail_titre": "Item-level detail", "aq_fiab_item_total": "Corrected item-total correlation", "aq_fiab_alpha_supprime": "Alpha if item deleted",
        "aq_fiab_normalite": "Composite score normality (Shapiro-Wilk)", "aq_fiab_normalite_oui": "normal distribution (p ≥ 0.05)",
        "aq_fiab_normalite_non": "non-normal distribution (p < 0.05)",
        "aq_correlations": "4. Correlations between constructs", "aq_corr_intro": "Pearson correlation between the composite scores of each construct, with significance testing.",
        "aq_afe_titre": "5. Exploratory factor analysis (EFA)",
        "aq_afe_kmo": "KMO index (sampling adequacy)", "aq_afe_bartlett": "Bartlett's test of sphericity",
        "aq_afe_facteurs": "Factors extracted (Kaiser criterion)", "aq_afe_variance": "Cumulative explained variance",
        "aq_afe_charges": "Factor loadings (after varimax rotation)",
        "aq_afc_titre": "6. Confirmatory factor analysis (CFA)",
        "aq_afc_intro": "Tests whether the postulated measurement structure (each construct measured by its declared items) fits the observed data well.",
        "aq_afc_indisponible": "The CFA could not be computed for this analysis (non-converging model or insufficient data) — the other results remain valid.",
        "aq_reg_titre": "7. Multiple regressions",
        "aq_reg_intro": "Effect of each predictor on the dependent variable, controlling simultaneously for the other predictors (standardized coefficients, directly comparable to each other).",
        "aq_reg_predicteur": "Predictor", "aq_reg_beta": "Standardized β", "aq_reg_p": "p", "aq_reg_ic": "95% CI",
        "aq_med_titre": "8. Mediation tests (bootstrap)",
        "aq_med_intro": "Indirect effect (a × b) re-estimated over 2000 resamples — mediation is considered significant when the 95% confidence interval excludes zero (Preacher & Hayes method).",
        "aq_med_chemin": "Path tested", "aq_med_indirect": "Indirect effect", "aq_med_ic": "95% CI", "aq_med_verdict": "Verdict",
        "aq_corr_v1": "Variable 1", "aq_corr_v2": "Variable 2", "aq_corr_r": "r", "aq_corr_p": "p", "aq_corr_methode": "Method",
        "aq_feuille_desc": "Descriptive statistics", "aq_feuille_freq": "Frequency tables",
        "aq_feuille_fiab": "Reliability (alpha)", "aq_feuille_corr": "Correlations",
        "aq_graphiques": "Charts", "aq_graph_fiabilite": "Construct reliability",
        "aq_graph_moyennes": "Mean scores by construct", "aq_graph_matrice": "Correlation matrix",
        "aq_synthese_titre": "9. Synthesis and interpretation", "aq_synthese_generale": "General synthesis",
        "aq_synthese_fiabilite": "Reliability discussion", "aq_synthese_hypotheses": "Research hypothesis testing",
        "aq_synthese_limites": "Limitations of the analysis", "aq_synthese_recommandations": "Recommendations",
        "aq_hyp_code": "Hypothesis", "aq_hyp_verdict": "Verdict", "aq_hyp_justif": "Justification",
        "aq_synthese_avertissement": "AI-generated interpretation based on the statistical results above — to be validated by the researcher. The correlations reported never formally establish mediation or moderation, only consistency or inconsistency with the tested hypothesis.",
        "aq_synthese_indisponible": "The interpretive synthesis could not be generated for this analysis — the raw statistical results above remain fully valid and usable.",
        "guide_grille_titre": "Coherence grid",
        "guide_grille_intro": "Each main question in the guide, mapped to the theoretical dimension it aims to explore — to be reviewed and amended by the researcher.",
        "guide_grille_question": "Question", "guide_grille_dimension": "Dimension targeted", "guide_grille_justif": "Rationale",
        "avertissement_ia": (
            "This document was automatically generated by artificial intelligence. It is a preparation "
            "aid tool — it must be reviewed, validated and adapted by the researcher before any use in "
            "the field or in academic work. It is not a scientifically validated instrument in the "
            "psychometric sense of the term."
        ),
    },
}


def _l(langue):
    return L.get(langue, L["fr"])


def _fmt_temps(secondes: float) -> str:
    h, reste = divmod(int(secondes or 0), 3600)
    m, s = divmod(reste, 60)
    return f"{h:d}:{m:02d}:{s:02d}" if h else f"{m:02d}:{s:02d}"


def _fmt_date(iso: str) -> str:
    if not iso:
        return "—"
    try:
        return datetime.fromisoformat(iso.replace("Z", "+00:00")).strftime("%d/%m/%Y")
    except ValueError:
        return iso[:10]


def _stats_segments(segments: list) -> dict:
    mots = [m for s in segments for m in (s.get("mots") or [])]
    duree = max((s.get("fin", 0) for s in segments), default=0)
    confiance_moy = (sum(m.get("confiance", 1) for m in mots) / len(mots)) if mots else None
    locuteurs = {}
    for s in segments:
        loc = s.get("locuteur")
        if loc:
            locuteurs[loc] = locuteurs.get(loc, 0) + max(0, s.get("fin", 0) - s.get("debut", 0))
    return {
        "duree_sec": duree, "nb_segments": len(segments), "nb_mots": len(mots),
        "confiance_moyenne": confiance_moy, "locuteurs": locuteurs,
    }


def _kappa_interpretation(k, langue="fr") -> str:
    termes = {
        "fr": ["accord faible", "accord passable", "accord modéré", "accord fort", "accord quasi parfait"],
        "en": ["poor agreement", "fair agreement", "moderate agreement", "substantial agreement", "almost perfect agreement"],
    }[langue if langue in ("fr", "en") else "fr"]
    if k is None:
        return "—"
    if k < 0.20:
        return termes[0]
    if k < 0.40:
        return termes[1]
    if k < 0.60:
        return termes[2]
    if k < 0.80:
        return termes[3]
    return termes[4]


def _stats_qualitatives(analyse: dict, entretiens_ordre: list) -> dict:
    """Indicateurs statistiques propres à la recherche qualitative :
    - fréquence par dimension (nb de thèmes/concepts/verbatims qui s'y rattachent) ;
    - convergence : pour chaque thème, sur combien d'entretiens distincts il apparaît ;
    - saturation théorique : nombre de concepts NOUVEAUX apportés par chaque entretien."""
    themes_par_nom = {t["theme"]: t for t in analyse.get("second_ordre", [])}
    concepts_par_nom = {c["concept"]: c for c in analyse.get("premier_ordre", [])}
    dimensions = analyse.get("dimensions_agregees") or []

    par_dimension = []
    for dim in dimensions:
        nb_themes = 0
        nb_concepts = 0
        nb_verbatims = 0
        for nom_t in dim.get("themes_lies", []):
            t = themes_par_nom.get(nom_t)
            if not t:
                continue
            nb_themes += 1
            for nom_c in t.get("concepts_lies", []):
                c = concepts_par_nom.get(nom_c)
                if not c:
                    continue
                nb_concepts += 1
                nb_verbatims += len(c.get("verbatims", []))
        par_dimension.append({
            "dimension": dim["dimension"], "nb_themes": nb_themes,
            "nb_concepts": nb_concepts, "nb_verbatims": nb_verbatims,
        })

    total_entretiens = len(entretiens_ordre) or 1
    convergence = []
    for t in analyse.get("second_ordre", []):
        entretiens_du_theme = set()
        for nom_c in t.get("concepts_lies", []):
            c = concepts_par_nom.get(nom_c)
            if not c:
                continue
            for v in c.get("verbatims", []):
                if v.get("entretien"):
                    entretiens_du_theme.add(v["entretien"])
        convergence.append({
            "theme": t["theme"], "nb_entretiens": len(entretiens_du_theme),
            "pct_entretiens": round(len(entretiens_du_theme) / total_entretiens * 100, 1),
        })
    convergence.sort(key=lambda x: -x["nb_entretiens"])

    saturation = []
    concepts_vus = set()
    for titre_e in entretiens_ordre:
        concepts_ici = set()
        for c in analyse.get("premier_ordre", []):
            for v in c.get("verbatims", []):
                if v.get("entretien") == titre_e:
                    concepts_ici.add(c["concept"])
        nouveaux = concepts_ici - concepts_vus
        concepts_vus |= concepts_ici
        saturation.append({
            "entretien": titre_e, "nouveaux_concepts": len(nouveaux),
            "cumul_concepts": len(concepts_vus),
        })

    return {"par_dimension": par_dimension, "convergence": convergence, "saturation": saturation}


# ----------------------------------------------------------------- utilitaires Word
def _champ_toc(document, langue):
    paragraphe = document.add_paragraph()
    run = paragraphe.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    texte_repli = OxmlElement("w:t")
    texte_repli.text = _l(langue)["sommaire_repli"]
    fld_sep.append(texte_repli)
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin); r.append(instr); r.append(fld_sep); r.append(fld_end)


def _numero_page(document, langue):
    footer = document.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _champ(instr_text):
        r = p.add_run()
        b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin"); r._r.append(b)
        i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = instr_text; r._r.append(i)
        s = OxmlElement("w:fldChar"); s.set(qn("w:fldCharType"), "separate"); r._r.append(s)
        e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end"); r._r.append(e)

    p.add_run(_l(langue)["page"])
    _champ("PAGE")
    p.add_run(_l(langue)["sur"])
    _champ("NUMPAGES")


def _avertissement_ia(document, langue):
    p = document.add_paragraph()
    run = p.add_run("⚠ " + _l(langue)["avertissement_ia"])
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x8A, 0x6A, 0x2E)
    document.add_paragraph()


def _page_de_garde(document, titre, sous_titre, lignes_meta):
    t = document.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(140)
    run = t.add_run(titre)
    run.font.size = Pt(30); run.font.bold = True; run.font.color.rgb = RGBColor(0x1B, 0x15, 0x03)

    st = document.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_after = Pt(60)
    run2 = st.add_run(sous_titre)
    run2.font.size = Pt(16); run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for ligne in lignes_meta:
        m = document.add_paragraph()
        m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        m.add_run(ligne).font.size = Pt(11)

    document.add_page_break()


def _ecrire_structure_analyse(document, analyse, langue):
    themes_par_nom = {t["theme"]: t for t in analyse.get("second_ordre", [])}
    concepts_par_nom = {c["concept"]: c for c in analyse.get("premier_ordre", [])}
    loc_prefixe = _l(langue)["locuteur_prefixe"]

    def _ecrire_theme(theme, niveau):
        document.add_heading(theme["theme"], level=niveau)
        if theme.get("description"):
            document.add_paragraph(theme["description"])
        for nom_c in theme.get("concepts_lies", []):
            c = concepts_par_nom.get(nom_c)
            if not c:
                continue
            document.add_paragraph(c["concept"], style="List Bullet")
            for v in c.get("verbatims", []):
                vp = document.add_paragraph(style="List Bullet 2")
                vp.add_run(f"« {v.get('texte', '')} » ").italic = True
                suffixe = f" — {v['entretien']}" if v.get("entretien") else ""
                vp.add_run(f"[{_fmt_temps(v.get('debut', 0))}{suffixe}]").font.size = Pt(9)

    dimensions = analyse.get("dimensions_agregees") or []
    if dimensions:
        for dim in dimensions:
            document.add_heading(dim["dimension"], level=2)
            for nom_t in dim.get("themes_lies", []):
                t = themes_par_nom.get(nom_t)
                if t:
                    _ecrire_theme(t, 3)
    else:
        for t in analyse.get("second_ordre", []):
            _ecrire_theme(t, 2)
    _ = loc_prefixe  # réservé pour cohérence future des libellés de locuteur dans le corps


def _ecrire_transcription(document, segments, langue):
    loc_prefixe = _l(langue)["locuteur_prefixe"]
    for s in segments:
        p = document.add_paragraph()
        tc = p.add_run(f"[{_fmt_temps(s.get('debut', 0))}] ")
        tc.bold = True
        tc.font.color.rgb = RGBColor(0xB0, 0x7A, 0x1E)
        if s.get("locuteur"):
            loc = p.add_run(f"{s['locuteur'].replace('SPEAKER_', loc_prefixe)} — ")
            loc.bold = True
        p.add_run(s.get("texte", ""))


# ----------------------------------------------------------------- Word — entretien seul
def generer_docx_entretien(entretien: dict) -> io.BytesIO:
    langue = entretien.get("analyse_langue") or "fr"
    l = _l(langue)
    doc = Document()
    _page_de_garde(
        doc, entretien.get("titre") or "Entretien",
        l["sous_titre_entretien"],
        [
            f"{l['langue_detectee']}{entretien.get('langue_detectee') or entretien.get('langue') or '—'}",
            f"{l['genere_le']}{datetime.now().strftime('%d/%m/%Y %H:%M')}",
        ],
    )
    _champ_toc(doc, langue)
    _avertissement_ia(doc, langue)
    doc.add_page_break()

    doc.add_heading(l["transcription"], level=1)
    _ecrire_transcription(doc, entretien.get("segments", []), langue)

    analyse = entretien.get("analyse")
    if analyse:
        doc.add_heading(l["analyse_qualitative"], level=1)
        methode = entretien.get("analyse_methode")
        doc.add_paragraph(
            f"{l['methode']}{METHODE_LABEL[langue].get(methode, methode or '—')}  ·  "
            f"{l['modele']}{entretien.get('analyse_modele') or '—'}"
        ).italic = True

        if analyse.get("demarche_methodologique"):
            doc.add_heading(f"2.1. {l['demarche_methodologique']}", level=2)
            doc.add_paragraph(analyse["demarche_methodologique"])

        doc.add_heading(f"2.2. {l['structure_resultats']}", level=2)
        _ecrire_structure_analyse(doc, analyse, langue)

        if analyse.get("synthese"):
            doc.add_heading(f"2.3. {l['synthese_interpretative']}", level=2)
            doc.add_paragraph(analyse["synthese"])
        if analyse.get("limites"):
            doc.add_heading(f"2.4. {l['limites_analyse']}", level=2)
            p = doc.add_paragraph(analyse["limites"])
            for run in p.runs:
                run.italic = True

        if methode and methode in REFERENCES_APA:
            doc.add_heading(l["reference"], level=2)
            doc.add_paragraph(REFERENCES_APA[methode])

    _numero_page(doc, langue)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ----------------------------------------------------------------- Word — étude (corpus)
def generer_docx_etude(corpus: dict, entretiens: list, fiabilite: dict = None) -> io.BytesIO:
    langue = corpus.get("analyse_langue") or "fr"
    l = _l(langue)
    doc = Document()
    analyse = corpus.get("analyse") or {}
    methode = corpus.get("analyse_methode")
    duree_totale = sum(_stats_segments(e.get("segments", []))["duree_sec"] for e in entretiens)
    langues = sorted({e.get("langue_detectee") or e.get("langue") or "—" for e in entretiens})

    _page_de_garde(
        doc, l["rapport_titre"], corpus.get("nom") or l["etude"],
        [
            f"{METHODE_LABEL[langue].get(methode, methode or '—')}",
            f"{len(entretiens)} {l['entretiens_mot']} · {_fmt_temps(duree_totale)}{l['materiau_audio']}",
            f"{l['langues_parentheses']}{', '.join(langues)}",
            f"{l['genere_le']}{datetime.now().strftime('%d/%m/%Y %H:%M')}",
        ],
    )
    _champ_toc(doc, langue)
    _avertissement_ia(doc, langue)
    doc.add_page_break()

    doc.add_heading(l["presentation_etude"], level=1)
    doc.add_paragraph(l["etude_intro"].format(n=len(entretiens), duree=_fmt_temps(duree_totale)))
    if corpus.get("analyse_contexte"):
        doc.add_paragraph(f"{l['question_recherche']}{corpus['analyse_contexte']}")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Light Grid Accent 1"
    for i, h in enumerate([l["entretien_col"], l["langue_col"], l["duree_col"]]):
        tbl.rows[0].cells[i].text = h
    for e in entretiens:
        row = tbl.add_row().cells
        row[0].text = e.get("titre") or "—"
        row[1].text = e.get("langue_detectee") or e.get("langue") or "—"
        row[2].text = _fmt_temps(_stats_segments(e.get("segments", []))["duree_sec"])

    doc.add_heading(l["demarche_titre2"], level=1)
    if analyse.get("demarche_methodologique"):
        doc.add_paragraph(analyse["demarche_methodologique"])
    else:
        doc.add_paragraph(l["aucune_analyse_corpus"])

    if analyse:
        doc.add_heading(l["structure_resultats3"], level=1)
        _ecrire_structure_analyse(doc, analyse, langue)

        doc.add_heading(l["synthese4"], level=1)
        doc.add_paragraph(analyse.get("synthese") or "—")

        doc.add_heading(l["limites5"], level=1)
        p = doc.add_paragraph(analyse.get("limites") or "—")
        for run in p.runs:
            run.italic = True

    if fiabilite and fiabilite.get("details"):
        doc.add_heading(l["fiabilite6"], level=1)
        doc.add_paragraph(l["fiabilite_intro"])
        tbl2 = doc.add_table(rows=1, cols=4)
        tbl2.style = "Light Grid Accent 1"
        for i, h in enumerate([l["entretien_col"], l["codeurs_col"], l["kappa_moyen_col"], l["interpretation_col"]]):
            tbl2.rows[0].cells[i].text = h
        for d in fiabilite["details"]:
            row = tbl2.add_row().cells
            row[0].text = d["titre"]
            row[1].text = str(d["nb_codeurs"])
            row[2].text = str(d["kappa_moyen"]) if d["kappa_moyen"] is not None else "—"
            row[3].text = _kappa_interpretation(d["kappa_moyen"], langue)

    if analyse:
        doc.add_heading(l["rigueur7"], level=1)
        entretiens_ordre = [e.get("titre", "") for e in entretiens]
        stats_q = _stats_qualitatives(analyse, entretiens_ordre)

        doc.add_heading(l["convergence71"], level=2)
        doc.add_paragraph(l["convergence_intro"])
        tbl3 = doc.add_table(rows=1, cols=3)
        tbl3.style = "Light Grid Accent 1"
        for i, h in enumerate([l["theme_col"], l["entretiens_concernes_col"], l["pct_corpus_col"]]):
            tbl3.rows[0].cells[i].text = h
        for c in stats_q["convergence"]:
            row = tbl3.add_row().cells
            row[0].text = c["theme"]
            row[1].text = str(c["nb_entretiens"])
            row[2].text = f"{c['pct_entretiens']} %"

        doc.add_heading(l["saturation72"], level=2)
        doc.add_paragraph(l["saturation_intro"])
        tbl4 = doc.add_table(rows=1, cols=3)
        tbl4.style = "Light Grid Accent 1"
        for i, h in enumerate([l["entretien_col"], l["concepts_nouveaux_col"], l["cumul_col"]]):
            tbl4.rows[0].cells[i].text = h
        for s in stats_q["saturation"]:
            row = tbl4.add_row().cells
            row[0].text = s["entretien"]
            row[1].text = str(s["nouveaux_concepts"])
            row[2].text = str(s["cumul_concepts"])

    if methode and methode in REFERENCES_APA:
        doc.add_heading(l["references"], level=1)
        doc.add_paragraph(REFERENCES_APA[methode])

    doc.add_page_break()
    doc.add_heading(l["annexe_a"], level=1)
    tblA = doc.add_table(rows=1, cols=4)
    tblA.style = "Light Grid Accent 1"
    for i, h in enumerate([l["titre_col"], l["langue_col"], l["duree_col"], l["date_col"]]):
        tblA.rows[0].cells[i].text = h
    for e in entretiens:
        row = tblA.add_row().cells
        row[0].text = e.get("titre") or "—"
        row[1].text = e.get("langue_detectee") or e.get("langue") or "—"
        row[2].text = _fmt_temps(_stats_segments(e.get("segments", []))["duree_sec"])
        row[3].text = _fmt_date(e.get("cree_le"))

    doc.add_page_break()
    doc.add_heading(l["annexe_b"], level=1)
    for e in entretiens:
        doc.add_heading(e.get("titre") or "—", level=2)
        _ecrire_transcription(doc, e.get("segments", []), langue)

    _numero_page(doc, langue)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ----------------------------------------------------------------- Excel
def _entete(ws, valeurs, ligne=1):
    for i, v in enumerate(valeurs, start=1):
        cell = ws.cell(row=ligne, column=i, value=v)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor=OR_HEX)
        cell.alignment = Alignment(vertical="center")
    for i in range(1, len(valeurs) + 1):
        ws.column_dimensions[get_column_letter(i)].width = 22


def _note_avertissement_xlsx(ws, ligne, langue):
    cell = ws.cell(ligne, 1, "⚠ " + _l(langue)["avertissement_ia"])
    cell.font = Font(italic=True, size=9, color="8A6A2E")
    cell.alignment = Alignment(wrap_text=True, vertical="top")
    ws.merge_cells(start_row=ligne, start_column=1, end_row=ligne, end_column=4)
    ws.row_dimensions[ligne].height = 45


def generer_xlsx_entretien(entretien: dict) -> io.BytesIO:
    langue = entretien.get("analyse_langue") or "fr"
    l = _l(langue)
    loc_prefixe = l["locuteur_prefixe"]
    wb = Workbook()

    ws_t = wb.active
    ws_t.title = l["feuille_transcription"]
    _entete(ws_t, [l["debut"], l["fin"], l["locuteur"], l["texte"], l["confiance_moy"]])
    for i, s in enumerate(entretien.get("segments", []), start=2):
        mots = s.get("mots") or []
        conf = round(sum(m.get("confiance", 1) for m in mots) / len(mots) * 100, 1) if mots else None
        ws_t.cell(i, 1, _fmt_temps(s.get("debut", 0)))
        ws_t.cell(i, 2, _fmt_temps(s.get("fin", 0)))
        ws_t.cell(i, 3, (s.get("locuteur") or "").replace("SPEAKER_", loc_prefixe))
        ws_t.cell(i, 4, s.get("texte", ""))
        ws_t.cell(i, 5, conf)
    ws_t.column_dimensions["D"].width = 70

    stats = _stats_segments(entretien.get("segments", []))
    ws_s = wb.create_sheet(l["feuille_statistiques"])
    _entete(ws_s, [l["indicateur"], l["valeur"]])
    lignes = [
        (l["titre_stat"], entretien.get("titre") or "—"),
        (l["langue_detectee_stat"], entretien.get("langue_detectee") or entretien.get("langue") or "—"),
        (l["duree_totale"], _fmt_temps(stats["duree_sec"])),
        (l["nb_segments"], stats["nb_segments"]),
        (l["nb_mots"], stats["nb_mots"]),
        (l["fiabilite_moy"], round(stats["confiance_moyenne"] * 100, 1) if stats["confiance_moyenne"] is not None else "—"),
    ]
    for i, (k, v) in enumerate(lignes, start=2):
        ws_s.cell(i, 1, k).font = Font(bold=True)
        ws_s.cell(i, 2, v)

    if stats["locuteurs"]:
        r = len(lignes) + 3
        ws_s.cell(r, 1, l["temps_parole"]).font = Font(bold=True)
        r += 1
        for loc, dur in sorted(stats["locuteurs"].items(), key=lambda x: -x[1]):
            ws_s.cell(r, 1, loc.replace("SPEAKER_", loc_prefixe))
            ws_s.cell(r, 2, _fmt_temps(dur))
            r += 1
        _note_avertissement_xlsx(ws_s, r + 1, langue)
    else:
        _note_avertissement_xlsx(ws_s, len(lignes) + 3, langue)

    analyse = entretien.get("analyse")
    if analyse:
        _feuille_analyse(wb, l["feuille_analyse"], analyse, langue)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def _feuille_analyse(wb, nom, analyse, langue):
    l = _l(langue)
    ws = wb.create_sheet(nom)
    _entete(ws, [l["niveau"], l["libelle"], l["detail_verbatim"], l["entretien_col"], l["horodatage"]])
    r = 2
    for c in analyse.get("premier_ordre", []):
        for v in c.get("verbatims", []) or [{}]:
            ws.cell(r, 1, l["premier_ordre"])
            ws.cell(r, 2, c.get("concept", ""))
            ws.cell(r, 3, v.get("texte", ""))
            ws.cell(r, 4, v.get("entretien", ""))
            ws.cell(r, 5, _fmt_temps(v.get("debut", 0)) if v else "")
            r += 1
    for t in analyse.get("second_ordre", []):
        ws.cell(r, 1, l["second_ordre"])
        ws.cell(r, 2, t.get("theme", ""))
        ws.cell(r, 3, t.get("description", ""))
        r += 1
    for d in analyse.get("dimensions_agregees", []) or []:
        ws.cell(r, 1, l["dimension_agregee"])
        ws.cell(r, 2, d.get("dimension", ""))
        ws.cell(r, 3, ", ".join(d.get("themes_lies", [])))
        r += 1
    ws.column_dimensions["C"].width = 60
    return ws


def _feuille_codebook(wb, analyse, langue):
    l = _l(langue)
    ws = wb.create_sheet(l["feuille_codebook"])
    _entete(ws, [l["dimension"], l["theme"], l["concept"], l["frequence"]])
    themes_par_nom = {t["theme"]: t for t in analyse.get("second_ordre", [])}
    concepts_par_nom = {c["concept"]: c for c in analyse.get("premier_ordre", [])}
    r = 2
    dimensions = analyse.get("dimensions_agregees") or []
    if dimensions:
        for dim in dimensions:
            for nom_t in dim.get("themes_lies", []):
                t = themes_par_nom.get(nom_t)
                if not t:
                    continue
                for nom_c in t.get("concepts_lies", []):
                    c = concepts_par_nom.get(nom_c)
                    if not c:
                        continue
                    ws.cell(r, 1, dim["dimension"])
                    ws.cell(r, 2, t["theme"])
                    ws.cell(r, 3, c["concept"])
                    ws.cell(r, 4, len(c.get("verbatims", [])))
                    r += 1
    else:
        for t in analyse.get("second_ordre", []):
            for nom_c in t.get("concepts_lies", []):
                c = concepts_par_nom.get(nom_c)
                if not c:
                    continue
                ws.cell(r, 1, "—")
                ws.cell(r, 2, t["theme"])
                ws.cell(r, 3, c["concept"])
                ws.cell(r, 4, len(c.get("verbatims", [])))
                r += 1
    ws.column_dimensions["B"].width = 30
    ws.column_dimensions["C"].width = 30


def _feuille_stats_qualitatives(wb, analyse, entretiens_ordre, langue):
    l = _l(langue)
    stats_q = _stats_qualitatives(analyse, entretiens_ordre)
    ws = wb.create_sheet(l["feuille_stats_qual"])

    ws.cell(1, 1, l["freq_dimension"]).font = Font(bold=True, size=13)
    _entete(ws, [l["dimension"], l["themes_pl"], l["concepts_pl"], l["verbatims_pl"]], ligne=2)
    r = 3
    for d in stats_q["par_dimension"]:
        ws.cell(r, 1, d["dimension"]); ws.cell(r, 2, d["nb_themes"])
        ws.cell(r, 3, d["nb_concepts"]); ws.cell(r, 4, d["nb_verbatims"])
        r += 1

    r += 1
    ws.cell(r, 1, l["convergence_titre"]).font = Font(bold=True, size=13)
    r += 1
    _entete(ws, [l["theme_col"], l["entretiens_concernes_col"], l["pct_corpus_col"]], ligne=r)
    r += 1
    for c in stats_q["convergence"]:
        ws.cell(r, 1, c["theme"]); ws.cell(r, 2, c["nb_entretiens"]); ws.cell(r, 3, f"{c['pct_entretiens']} %")
        r += 1

    r += 1
    ws.cell(r, 1, l["saturation_titre"]).font = Font(bold=True, size=13)
    r += 1
    _entete(ws, [l["entretien_col"], l["concepts_nouveaux_col"], l["concepts_distincts"]], ligne=r)
    r += 1
    for s in stats_q["saturation"]:
        ws.cell(r, 1, s["entretien"]); ws.cell(r, 2, s["nouveaux_concepts"]); ws.cell(r, 3, s["cumul_concepts"])
        r += 1

    ws.column_dimensions["A"].width = 34


def generer_xlsx_etude(corpus: dict, entretiens: list, fiabilite: dict = None) -> io.BytesIO:
    langue = corpus.get("analyse_langue") or "fr"
    l = _l(langue)
    loc_prefixe = l["locuteur_prefixe"]
    wb = Workbook()
    analyse = corpus.get("analyse") or {}
    duree_totale = sum(_stats_segments(e.get("segments", []))["duree_sec"] for e in entretiens)
    tous_mots = sum(_stats_segments(e.get("segments", []))["nb_mots"] for e in entretiens)
    confiances = [
        st["confiance_moyenne"] for st in
        (_stats_segments(e.get("segments", [])) for e in entretiens)
        if st["confiance_moyenne"] is not None
    ]

    ws_v = wb.active
    ws_v.title = l["feuille_vue_ensemble"]
    _entete(ws_v, [l["indicateur"], l["valeur"]])
    lignes = [
        (l["etude_corpus"], corpus.get("nom") or "—"),
        (l["methodologie"], METHODE_LABEL[langue].get(corpus.get("analyse_methode"), corpus.get("analyse_methode") or "—")),
        (l["question_recherche_stat"], corpus.get("analyse_contexte") or "—"),
        (l["nb_entretiens"], len(entretiens)),
        (l["duree_cumulee"], _fmt_temps(duree_totale)),
        (l["nb_mots_total"], tous_mots),
        (l["fiabilite_moy_corpus"], round(sum(confiances) / len(confiances) * 100, 1) if confiances else "—"),
        (l["rapport_genere_le"], datetime.now().strftime("%d/%m/%Y %H:%M")),
    ]
    for i, (k, v) in enumerate(lignes, start=2):
        ws_v.cell(i, 1, k).font = Font(bold=True)
        ws_v.cell(i, 2, v)
    ws_v.column_dimensions["B"].width = 60
    _note_avertissement_xlsx(ws_v, len(lignes) + 3, langue)

    ws_e = wb.create_sheet(l["feuille_entretiens"])
    _entete(ws_e, [l["titre_col"], l["langue_col"], l["duree_col"], l["segments"], l["mots"], l["confiance_pct"], l["date"]])
    for i, e in enumerate(entretiens, start=2):
        st = _stats_segments(e.get("segments", []))
        ws_e.cell(i, 1, e.get("titre", ""))
        ws_e.cell(i, 2, e.get("langue_detectee") or e.get("langue") or "—")
        ws_e.cell(i, 3, _fmt_temps(st["duree_sec"]))
        ws_e.cell(i, 4, st["nb_segments"])
        ws_e.cell(i, 5, st["nb_mots"])
        ws_e.cell(i, 6, round(st["confiance_moyenne"] * 100, 1) if st["confiance_moyenne"] is not None else "—")
        ws_e.cell(i, 7, _fmt_date(e.get("cree_le")))

    ws_tr = wb.create_sheet(l["feuille_transcriptions"])
    _entete(ws_tr, [l["entretien_col"], l["debut"], l["fin"], l["locuteur"], l["texte"]])
    r = 2
    for e in entretiens:
        for s in e.get("segments", []):
            ws_tr.cell(r, 1, e.get("titre", ""))
            ws_tr.cell(r, 2, _fmt_temps(s.get("debut", 0)))
            ws_tr.cell(r, 3, _fmt_temps(s.get("fin", 0)))
            ws_tr.cell(r, 4, (s.get("locuteur") or "").replace("SPEAKER_", loc_prefixe))
            ws_tr.cell(r, 5, s.get("texte", ""))
            r += 1
    ws_tr.column_dimensions["E"].width = 70

    if analyse:
        _feuille_analyse(wb, l["feuille_analyse_transversale"], analyse, langue)
        _feuille_codebook(wb, analyse, langue)
        _feuille_stats_qualitatives(wb, analyse, [e.get("titre", "") for e in entretiens], langue)

    if fiabilite and fiabilite.get("details"):
        ws_f = wb.create_sheet(l["feuille_fiabilite"])
        _entete(ws_f, [l["entretien_col"], l["codeurs_col"], l["kappa_moyen_col"], l["interpretation_col"]])
        for i, d in enumerate(fiabilite["details"], start=2):
            ws_f.cell(i, 1, d["titre"])
            ws_f.cell(i, 2, d["nb_codeurs"])
            ws_f.cell(i, 3, d["kappa_moyen"])
            ws_f.cell(i, 4, _kappa_interpretation(d["kappa_moyen"], langue))

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ----------------------------------------------------------------- Word — guide d'entretien
def generer_docx_guide(guide_data: dict) -> io.BytesIO:
    langue = guide_data.get("langue") or "fr"
    l = _l(langue)
    g = guide_data.get("guide") or {}
    doc = Document()

    _page_de_garde(
        doc, g.get("titre") or guide_data.get("theme") or "—", l["guide_sous_titre"],
        [
            f"{l['guide_theme']}{guide_data.get('theme', '—')}",
            f"{l['guide_question']}{guide_data.get('question_recherche') or '—'}",
            f"{l['genere_le']}{datetime.now().strftime('%d/%m/%Y %H:%M')}",
        ],
    )
    _avertissement_ia(doc, langue)

    infos = g.get("informations_pratiques") or {}
    if infos:
        doc.add_heading(l["guide_infos_pratiques"], level=1)
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Light Grid Accent 1"
        paires = [
            (l["guide_type"], infos.get("type_entretien")),
            (l["guide_duree"], infos.get("duree_estimee")),
            (l["guide_population"], infos.get("population_cible")),
            (l["guide_materiel"], infos.get("materiel_recommande")),
        ]
        for label, valeur in paires:
            if not valeur:
                continue
            row = tbl.add_row().cells
            row[0].text = label
            row[0].paragraphs[0].runs[0].bold = True
            row[1].text = valeur

    if g.get("preambule"):
        doc.add_heading(l["guide_preambule"], level=1)
        p = doc.add_paragraph(g["preambule"])
        for run in p.runs:
            run.italic = True

    for i, section in enumerate(g.get("sections", []), start=1):
        doc.add_heading(f"{i}. {section.get('titre', '')}", level=1)
        if section.get("objectif"):
            obj = doc.add_paragraph()
            obj.add_run(l["guide_objectif"]).bold = True
            obj.add_run(section["objectif"])
        for q in section.get("questions", []):
            doc.add_paragraph(q.get("question", ""), style="List Bullet")
            relances = q.get("relances") or []
            if relances:
                rp = doc.add_paragraph(style="List Bullet 2")
                rp.add_run(l["guide_relances"]).italic = True
                for rel in relances:
                    doc.add_paragraph(rel, style="List Bullet 2")

    if g.get("grille_coherence"):
        doc.add_heading(l["guide_grille_titre"], level=1)
        doc.add_paragraph(l["guide_grille_intro"])
        tblg = doc.add_table(rows=1, cols=3)
        tblg.style = "Light Grid Accent 1"
        for i, h in enumerate([l["guide_grille_question"], l["guide_grille_dimension"], l["guide_grille_justif"]]):
            tblg.rows[0].cells[i].text = h
        for item in g["grille_coherence"]:
            row = tblg.add_row().cells
            row[0].text = item.get("question", "")
            row[1].text = item.get("dimension_visee", "")
            row[2].text = item.get("justification", "")

    if g.get("conseils_methodologiques"):
        doc.add_heading(l["guide_conseils"], level=1)
        doc.add_paragraph(g["conseils_methodologiques"])

    if g.get("note_methodologique"):
        doc.add_heading(l["guide_note"], level=1)
        p = doc.add_paragraph(g["note_methodologique"])
        for run in p.runs:
            run.italic = True

    _numero_page(doc, langue)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ----------------------------------------------------------------- Word — étude quantitative
def generer_docx_etude_quant(etude_data: dict) -> io.BytesIO:
    langue = etude_data.get("langue") or "fr"
    l = _l(langue)
    c = etude_data.get("contenu") or {}
    doc = Document()

    _page_de_garde(
        doc, c.get("titre") or etude_data.get("theme") or "—", l["eq_sous_titre"],
        [
            f"{l['eq_theme']}{etude_data.get('theme', '—')}",
            f"{l['eq_question']}{etude_data.get('question_recherche') or '—'}",
            f"{l['genere_le']}{datetime.now().strftime('%d/%m/%Y %H:%M')}",
        ],
    )
    _champ_toc(doc, langue)
    _avertissement_ia(doc, langue)
    doc.add_page_break()

    if c.get("cadre_theorique"):
        doc.add_heading(l["eq_cadre"], level=1)
        doc.add_paragraph(c["cadre_theorique"])

    if c.get("revue_litterature"):
        doc.add_heading(l["eq_revue"], level=1)
        doc.add_paragraph(c["revue_litterature"])

    m = c.get("methodologie") or {}
    if m:
        doc.add_heading(l["eq_methodo"], level=1)
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Light Grid Accent 1"
        for label, cle in [(l["eq_type_etude"], "type_etude"), (l["eq_population"], "population_cible"), (l["eq_echantillon"], "echantillon")]:
            if m.get(cle):
                row = tbl.add_row().cells
                row[0].text = label
                row[0].paragraphs[0].runs[0].bold = True
                row[1].text = m[cle]

        if m.get("hypotheses"):
            doc.add_heading(l["eq_hypotheses"], level=2)
            tblh = doc.add_table(rows=1, cols=2)
            tblh.style = "Light Grid Accent 1"
            tblh.rows[0].cells[0].text = l["eq_hyp_col"]
            tblh.rows[0].cells[1].text = l["eq_hyp_enonce"]
            for h in m["hypotheses"]:
                row = tblh.add_row().cells
                row[0].text = h.get("code", "")
                row[1].text = h.get("enonce", "")

        if m.get("variables"):
            doc.add_heading(l["eq_variables"], level=2)
            tblv = doc.add_table(rows=1, cols=3)
            tblv.style = "Light Grid Accent 1"
            for i, hh in enumerate([l["eq_var_nom"], l["eq_var_type"], l["eq_var_def"]]):
                tblv.rows[0].cells[i].text = hh
            for v in m["variables"]:
                row = tblv.add_row().cells
                row[0].text = v.get("nom", "")
                row[1].text = v.get("type", "")
                row[2].text = v.get("definition", "")

    q = c.get("questionnaire") or {}
    if q.get("sections"):
        doc.add_heading(l["eq_questionnaire"], level=1)
        for section in q["sections"]:
            doc.add_heading(section.get("titre", ""), level=2)
            for item in section.get("items", []):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"[{item.get('code', '')}] ").bold = True
                p.add_run(item.get("libelle", ""))
                if item.get("options"):
                    doc.add_paragraph(" / ".join(item["options"]), style="List Bullet 2")

    if c.get("note_methodologique"):
        doc.add_heading(l["eq_note"], level=1)
        p = doc.add_paragraph(c["note_methodologique"])
        for run in p.runs:
            run.italic = True

    refs = c.get("references_apa") or {}
    if refs.get("methodologie") or refs.get("concepts_a_referencer"):
        doc.add_heading(l["eq_references"], level=1)
        if refs.get("methodologie"):
            p_intro_m = doc.add_paragraph()
            p_intro_m.add_run(l["eq_ref_methodo_intro"]).bold = True
            for ref in refs["methodologie"]:
                doc.add_paragraph(ref, style="List Bullet")
        if refs.get("concepts_a_referencer"):
            p_intro = doc.add_paragraph()
            p_intro.add_run(l["eq_ref_concepts_intro"]).bold = True
            for concept in refs["concepts_a_referencer"]:
                doc.add_paragraph(f"{concept.get('concept', '')} — {concept.get('auteur_associe', '')}", style="List Bullet")

    _numero_page(doc, langue)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ----------------------------------------------------------------- Excel — gabarit du questionnaire
def generer_xlsx_template_questionnaire(etude_data: dict) -> io.BytesIO:
    langue = etude_data.get("langue") or "fr"
    l = _l(langue)
    q = (etude_data.get("contenu") or {}).get("questionnaire") or {}
    items = [item for section in q.get("sections", []) for item in section.get("items", [])]

    wb = Workbook()
    ws_r = wb.active
    ws_r.title = l["eq_feuille_reponses"]
    _entete(ws_r, [item["code"] for item in items])
    for i in range(1, len(items) + 1):
        ws_r.column_dimensions[get_column_letter(i)].width = 16

    ws_g = wb.create_sheet(l["eq_feuille_guide"])
    ws_g.cell(1, 1, l["eq_guide_intro"]).font = Font(italic=True)
    ws_g.merge_cells(start_row=1, start_column=1, end_row=1, end_column=4)
    ws_g.row_dimensions[1].height = 45
    _entete(ws_g, [l["eq_guide_code"], l["eq_guide_libelle"], l["eq_guide_type"], l["eq_guide_valeurs"]], ligne=2)
    for i, item in enumerate(items, start=3):
        ws_g.cell(i, 1, item["code"])
        ws_g.cell(i, 2, item.get("libelle", ""))
        ws_g.cell(i, 3, l["eq_type_libelle"].get(item.get("type", ""), item.get("type", "")))
        if item.get("type") == "echelle_likert":
            valeurs = f"{item.get('echelle_min', 1)} à {item.get('echelle_max', 5)}"
            if item.get("options"):
                valeurs += " (" + " / ".join(item["options"]) + ")"
        else:
            valeurs = " / ".join(item.get("options", [])) if item.get("options") else "—"
        ws_g.cell(i, 4, valeurs)
    ws_g.column_dimensions["B"].width = 45
    ws_g.column_dimensions["D"].width = 50

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ----------------------------------------------------------------- Word — analyse quantitative
def generer_docx_analyse_quant(etude_data: dict, analyse_data: dict) -> io.BytesIO:
    langue = etude_data.get("langue") or "fr"
    l = _l(langue)
    r = analyse_data.get("resultats") or {}
    c = etude_data.get("contenu") or {}
    doc = Document()

    _page_de_garde(
        doc, c.get("titre") or etude_data.get("theme") or "—", l["aq_sous_titre"],
        [
            f"{l['aq_n_repondants']} : {r.get('n_repondants', '—')}",
            f"{l['genere_le']}{datetime.now().strftime('%d/%m/%Y %H:%M')}",
        ],
    )
    _avertissement_ia(doc, langue)
    doc.add_page_break()

    if r.get("descriptives"):
        doc.add_heading(l["aq_descriptives"], level=1)
        tbl = doc.add_table(rows=1, cols=9)
        tbl.style = "Light Grid Accent 1"
        entetes = [l["aq_desc_col"], l["aq_desc_n"], l["aq_desc_moy"], l["aq_desc_ic95"], l["aq_desc_mediane"],
                   l["aq_desc_et"], l["aq_desc_asymetrie"], l["aq_desc_aplatissement"], f"{l['aq_desc_min']}/{l['aq_desc_max']}"]
        for i, h in enumerate(entetes):
            tbl.rows[0].cells[i].text = h
        for d in r["descriptives"]:
            row = tbl.add_row().cells
            row[0].text = f"{d['code']} — {d['libelle']}"
            row[1].text = str(d["n"])
            row[2].text = str(d["moyenne"])
            row[3].text = f"[{d['ic95_bas']} ; {d['ic95_haut']}]" if d.get("ic95_bas") is not None else "—"
            row[4].text = str(d["mediane"])
            row[5].text = str(d["ecart_type"]) if d["ecart_type"] is not None else "—"
            row[6].text = str(d["asymetrie"]) if d.get("asymetrie") is not None else "—"
            row[7].text = str(d["aplatissement"]) if d.get("aplatissement") is not None else "—"
            row[8].text = f"{d['min']} / {d['max']}"

    if r.get("frequences"):
        doc.add_heading(l["aq_frequences"], level=1)
        for f in r["frequences"]:
            doc.add_heading(f"{f['code']} — {f['libelle']}", level=2)
            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = "Light Grid Accent 1"
            for i, h in enumerate([l["aq_freq_modalite"], l["aq_freq_effectif"], l["aq_freq_pct"]]):
                tbl.rows[0].cells[i].text = h
            for m in f["modalites"]:
                row = tbl.add_row().cells
                row[0].text = str(m["valeur"])
                row[1].text = str(m["effectif"])
                row[2].text = f"{m['pourcentage']} %"

    if r.get("fiabilite"):
        doc.add_heading(l["aq_fiabilite"], level=1)
        doc.add_paragraph(l["aq_fiab_intro"])
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Light Grid Accent 1"
        for i, h in enumerate([l["aq_fiab_variable"], l["aq_fiab_nb_items"], l["aq_fiab_alpha"], l["aq_fiab_interpretation"]]):
            tbl.rows[0].cells[i].text = h
        for f in r["fiabilite"]:
            row = tbl.add_row().cells
            row[0].text = f["variable"]
            row[1].text = str(f["nb_items"])
            row[2].text = str(f["alpha_cronbach"]) if f["alpha_cronbach"] is not None else "—"
            row[3].text = f["interpretation"]

        for f in r["fiabilite"]:
            doc.add_heading(f"{f['variable']} — {l['aq_fiab_detail_titre']}", level=2)
            norm = f.get("normalite_score") or {}
            if norm.get("p_valeur") is not None:
                texte_norm = l["aq_fiab_normalite_oui"] if norm.get("distribution_normale") else l["aq_fiab_normalite_non"]
                doc.add_paragraph(f"{l['aq_fiab_normalite']} : W = {norm['statistique_shapiro']}, p = {norm['p_valeur']} — {texte_norm}").italic = True
            tbld = doc.add_table(rows=1, cols=3)
            tbld.style = "Light Grid Accent 1"
            for i, h in enumerate([l["aq_desc_col"], l["aq_fiab_item_total"], l["aq_fiab_alpha_supprime"]]):
                tbld.rows[0].cells[i].text = h
            for di in f.get("detail_items", []):
                row = tbld.add_row().cells
                row[0].text = di["code"]
                row[1].text = str(di["correlation_item_total"]) if di["correlation_item_total"] is not None else "—"
                row[2].text = str(di["alpha_si_supprime"]) if di["alpha_si_supprime"] is not None else "—"

    if r.get("correlations"):
        doc.add_heading(l["aq_correlations"], level=1)
        doc.add_paragraph(l["aq_corr_intro"])
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Light Grid Accent 1"
        for i, h in enumerate([l["aq_corr_v1"], l["aq_corr_v2"], l["aq_corr_methode"], l["aq_corr_r"], l["aq_corr_p"], l["interpretation_col"]]):
            tbl.rows[0].cells[i].text = h
        for cr in r["correlations"]:
            row = tbl.add_row().cells
            row[0].text = cr["variable_1"]
            row[1].text = cr["variable_2"]
            row[2].text = cr.get("methode", "Pearson")
            row[3].text = str(cr["r"])
            row[4].text = str(cr["p_valeur"])
            row[5].text = cr["interpretation"]

    # --- Graphiques ---
    if r.get("fiabilite"):
        doc.add_heading(l["aq_graphiques"], level=1)
        for image_bytes, legende in [
            (graphique_fiabilite(r["fiabilite"]), l["aq_graph_fiabilite"]),
            (graphique_moyennes(r["fiabilite"]), l["aq_graph_moyennes"]),
            (graphique_matrice_correlations(r["fiabilite"], r.get("correlations") or []), l["aq_graph_matrice"]),
        ]:
            if image_bytes is None:
                continue
            doc.add_picture(io.BytesIO(image_bytes), width=Inches(6.2))
            legende_p = doc.add_paragraph(legende)
            legende_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in legende_p.runs:
                run.italic = True
                run.font.size = Pt(9)

    # --- Analyses avancées : AFE, AFC, régressions, médiations ---
    av = r.get("analyses_avancees") or {}

    if av.get("afe"):
        a = av["afe"]
        doc.add_heading(l["aq_afe_titre"], level=1)
        doc.add_paragraph(f"{l['aq_afe_kmo']} : {a['kmo_total']} ({a['kmo_interpretation']})")
        doc.add_paragraph(
            f"{l['aq_afe_bartlett']} : χ² = {a['bartlett_chi2']}, p = {a['bartlett_p']} "
            f"({'factorisable' if a['bartlett_factorisable'] else 'non factorisable'})"
        )
        doc.add_paragraph(f"{l['aq_afe_facteurs']} : {a['n_facteurs_extraits']}")
        doc.add_paragraph(f"{l['aq_afe_variance']} : {a['variance_expliquee_cumulee_pct']}%")
        doc.add_paragraph(l["aq_afe_charges"]).runs[0].bold = True
        tbl_afe = doc.add_table(rows=1, cols=1 + a["n_facteurs_extraits"])
        tbl_afe.style = "Light Grid Accent 1"
        tbl_afe.rows[0].cells[0].text = "Item"
        for fi in range(a["n_facteurs_extraits"]):
            tbl_afe.rows[0].cells[fi + 1].text = f"Facteur {fi + 1}"
        for c in a["charges_factorielles"]:
            row = tbl_afe.add_row().cells
            row[0].text = c["item"]
            for fi, charge in enumerate(c["charges"]):
                row[fi + 1].text = str(charge)

    if av.get("afc"):
        doc.add_heading(l["aq_afc_titre"], level=1)
        doc.add_paragraph(l["aq_afc_intro"])
        if av["afc"].get("erreur"):
            doc.add_paragraph(l["aq_afc_indisponible"])
        else:
            c = av["afc"]
            doc.add_paragraph(f"CFI = {c['cfi']}  |  TLI = {c['tli']}  |  RMSEA = {c['rmsea']}  |  SRMR = {c['srmr']}")
            p_interp = doc.add_paragraph(c["interpretation"])
            p_interp.runs[0].italic = True
            if c.get("charges_factorielles"):
                tbl_afc = doc.add_table(rows=1, cols=4)
                tbl_afc.style = "Light Grid Accent 1"
                for i, h in enumerate(["Construit", "Item", "Charge", "p"]):
                    tbl_afc.rows[0].cells[i].text = h
                for ch in c["charges_factorielles"]:
                    row = tbl_afc.add_row().cells
                    row[0].text = ch["construit"]; row[1].text = ch["item"]
                    row[2].text = str(ch["charge"]); row[3].text = str(ch["p_valeur"]) if ch["p_valeur"] is not None else "—"

    if av.get("regressions"):
        doc.add_heading(l["aq_reg_titre"], level=1)
        doc.add_paragraph(l["aq_reg_intro"])
        for r in av["regressions"]:
            if r.get("erreur"):
                continue
            p_titre = doc.add_paragraph(f"Variable dépendante : {r['dependante']}  (R² = {r['r2']}, R² ajusté = {r['r2_ajuste']}, F = {r['f_stat']}, p = {r['f_p']})")
            p_titre.runs[0].bold = True
            tbl_reg = doc.add_table(rows=1, cols=4)
            tbl_reg.style = "Light Grid Accent 1"
            for i, h in enumerate([l["aq_reg_predicteur"], l["aq_reg_beta"], l["aq_reg_p"], l["aq_reg_ic"]]):
                tbl_reg.rows[0].cells[i].text = h
            for p in r["predicteurs"]:
                row = tbl_reg.add_row().cells
                row[0].text = p["nom"] + (" *" if p["significatif"] else "")
                row[1].text = str(p["beta"]); row[2].text = str(p["p_valeur"])
                row[3].text = f"[{p['ic95_bas']} ; {p['ic95_haut']}]"

    if av.get("mediations"):
        doc.add_heading(l["aq_med_titre"], level=1)
        doc.add_paragraph(l["aq_med_intro"])
        tbl_med = doc.add_table(rows=1, cols=4)
        tbl_med.style = "Light Grid Accent 1"
        for i, h in enumerate([l["aq_med_chemin"], l["aq_med_indirect"], l["aq_med_ic"], l["aq_med_verdict"]]):
            tbl_med.rows[0].cells[i].text = h
        for med in av["mediations"]:
            if med.get("erreur"):
                continue
            row = tbl_med.add_row().cells
            row[0].text = f"{med['independante']} → {med['mediatrice']} → {med['dependante']}"
            row[1].text = str(med["effet_indirect_a_x_b"])
            row[2].text = f"[{med['ic95_bas']} ; {med['ic95_haut']}]"
            row[3].text = med["type_mediation"] if med["mediation_significative"] else "non significative"

    # --- Synthèse interprétative générée par IA ---
    synthese = analyse_data.get("synthese_interpretative")
    if synthese:
        doc.add_page_break()
        doc.add_heading(l["aq_synthese_titre"], level=1)
        p_avert = doc.add_paragraph(l["aq_synthese_avertissement"])
        for run in p_avert.runs:
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x8A, 0x85, 0x74)

        if synthese.get("synthese_generale"):
            doc.add_heading(l["aq_synthese_generale"], level=2)
            doc.add_paragraph(synthese["synthese_generale"])

        if synthese.get("fiabilite_discussion"):
            doc.add_heading(l["aq_synthese_fiabilite"], level=2)
            doc.add_paragraph(synthese["fiabilite_discussion"])

        if synthese.get("tests_hypotheses"):
            doc.add_heading(l["aq_synthese_hypotheses"], level=2)
            tblh = doc.add_table(rows=1, cols=3)
            tblh.style = "Light Grid Accent 1"
            for i, h in enumerate([l["aq_hyp_code"], l["aq_hyp_verdict"], l["aq_hyp_justif"]]):
                tblh.rows[0].cells[i].text = h
            for th in synthese["tests_hypotheses"]:
                row = tblh.add_row().cells
                row[0].text = th.get("code", "")
                row[1].text = th.get("verdict", "")
                row[2].text = th.get("justification", "")

        if synthese.get("limites"):
            doc.add_heading(l["aq_synthese_limites"], level=2)
            doc.add_paragraph(synthese["limites"])

        if synthese.get("recommandations"):
            doc.add_heading(l["aq_synthese_recommandations"], level=2)
            doc.add_paragraph(synthese["recommandations"])
    elif analyse_data.get("synthese_statut") == "non_disponible":
        doc.add_heading(l["aq_synthese_titre"], level=1)
        doc.add_paragraph(l["aq_synthese_indisponible"])

    _numero_page(doc, langue)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ----------------------------------------------------------------- Excel — analyse quantitative
def generer_xlsx_analyse_quant(etude_data: dict, analyse_data: dict) -> io.BytesIO:
    langue = etude_data.get("langue") or "fr"
    l = _l(langue)
    r = analyse_data.get("resultats") or {}
    wb = Workbook()

    ws_v = wb.active
    ws_v.title = l["feuille_vue_ensemble"]
    _entete(ws_v, [l["indicateur"], l["valeur"]])
    ws_v.cell(2, 1, l["aq_n_repondants"]).font = Font(bold=True)
    ws_v.cell(2, 2, r.get("n_repondants", "—"))
    ws_v.cell(3, 1, l["genere_le"]).font = Font(bold=True)
    ws_v.cell(3, 2, datetime.now().strftime("%d/%m/%Y %H:%M"))
    ws_v.column_dimensions["B"].width = 40
    _note_avertissement_xlsx(ws_v, 5, langue)

    if r.get("descriptives"):
        ws_d = wb.create_sheet(l["aq_feuille_desc"])
        _entete(ws_d, [l["aq_desc_col"], l["aq_desc_n"], l["aq_desc_moy"], l["aq_desc_ic95"], l["aq_desc_mediane"],
                       l["aq_desc_et"], l["aq_desc_asymetrie"], l["aq_desc_aplatissement"], l["aq_desc_min"], l["aq_desc_max"]])
        for i, d in enumerate(r["descriptives"], start=2):
            ws_d.cell(i, 1, f"{d['code']} — {d['libelle']}")
            ws_d.cell(i, 2, d["n"])
            ws_d.cell(i, 3, d["moyenne"])
            ws_d.cell(i, 4, f"[{d['ic95_bas']} ; {d['ic95_haut']}]" if d.get("ic95_bas") is not None else "—")
            ws_d.cell(i, 5, d["mediane"])
            ws_d.cell(i, 6, d["ecart_type"])
            ws_d.cell(i, 7, d.get("asymetrie"))
            ws_d.cell(i, 8, d.get("aplatissement"))
            ws_d.cell(i, 9, d["min"])
            ws_d.cell(i, 10, d["max"])
        ws_d.column_dimensions["A"].width = 40

    if r.get("frequences"):
        ws_f = wb.create_sheet(l["aq_feuille_freq"])
        _entete(ws_f, [l["aq_desc_col"], l["aq_freq_modalite"], l["aq_freq_effectif"], l["aq_freq_pct"]])
        row = 2
        for f in r["frequences"]:
            for m in f["modalites"]:
                ws_f.cell(row, 1, f"{f['code']} — {f['libelle']}")
                ws_f.cell(row, 2, m["valeur"])
                ws_f.cell(row, 3, m["effectif"])
                ws_f.cell(row, 4, m["pourcentage"])
                row += 1
        ws_f.column_dimensions["A"].width = 40

    if r.get("fiabilite"):
        ws_r = wb.create_sheet(l["aq_feuille_fiab"])
        _entete(ws_r, [l["aq_fiab_variable"], l["aq_fiab_nb_items"], l["aq_fiab_alpha"], l["aq_fiab_interpretation"],
                       l["aq_desc_moy"], l["aq_desc_et"], l["aq_fiab_normalite"]])
        for i, f in enumerate(r["fiabilite"], start=2):
            ws_r.cell(i, 1, f["variable"])
            ws_r.cell(i, 2, f["nb_items"])
            ws_r.cell(i, 3, f["alpha_cronbach"])
            ws_r.cell(i, 4, f["interpretation"])
            ws_r.cell(i, 5, f.get("moyenne_composite"))
            ws_r.cell(i, 6, f.get("ecart_type_composite"))
            norm = f.get("normalite_score") or {}
            if norm.get("p_valeur") is not None:
                texte_norm = l["aq_fiab_normalite_oui"] if norm.get("distribution_normale") else l["aq_fiab_normalite_non"]
                ws_r.cell(i, 7, f"W={norm['statistique_shapiro']}, p={norm['p_valeur']} — {texte_norm}")
        ws_r.column_dimensions["A"].width = 30
        ws_r.column_dimensions["G"].width = 45
        nb_construits = len(r["fiabilite"])

        # Graphique natif Excel (lié aux cellules, pas une image figée) : alpha par construit
        graphique = BarChart()
        graphique.type = "bar"
        graphique.title = l["aq_graph_fiabilite"]
        graphique.y_axis.title = l["aq_fiab_variable"]
        graphique.x_axis.title = l["aq_fiab_alpha"]
        donnees_alpha = Reference(ws_r, min_col=3, min_row=1, max_row=nb_construits + 1)
        categories = Reference(ws_r, min_col=1, min_row=2, max_row=nb_construits + 1)
        graphique.add_data(donnees_alpha, titles_from_data=True)
        graphique.set_categories(categories)
        graphique.height, graphique.width = 8, 16
        ws_r.add_chart(graphique, "I2")

        # Second graphique natif : moyennes composites par construit
        graphique_moy = BarChart()
        graphique_moy.type = "col"
        graphique_moy.title = l["aq_graph_moyennes"]
        graphique_moy.y_axis.title = l["aq_desc_moy"]
        donnees_moy = Reference(ws_r, min_col=5, min_row=1, max_row=nb_construits + 1)
        graphique_moy.add_data(donnees_moy, titles_from_data=True)
        graphique_moy.set_categories(categories)
        graphique_moy.height, graphique_moy.width = 8, 16
        ws_r.add_chart(graphique_moy, "I18")

        ws_ri = wb.create_sheet(l["aq_fiab_detail_titre"])
        _entete(ws_ri, [l["aq_fiab_variable"], l["aq_desc_col"], l["aq_fiab_item_total"], l["aq_fiab_alpha_supprime"]])
        row = 2
        for f in r["fiabilite"]:
            for di in f.get("detail_items", []):
                ws_ri.cell(row, 1, f["variable"])
                ws_ri.cell(row, 2, di["code"])
                ws_ri.cell(row, 3, di["correlation_item_total"])
                ws_ri.cell(row, 4, di["alpha_si_supprime"])
                row += 1
        ws_ri.column_dimensions["A"].width = 30

    if r.get("correlations"):
        ws_c = wb.create_sheet(l["aq_feuille_corr"])
        _entete(ws_c, [l["aq_corr_v1"], l["aq_corr_v2"], l["aq_corr_methode"], l["aq_corr_r"], l["aq_corr_p"], l["interpretation_col"]])
        for i, cr in enumerate(r["correlations"], start=2):
            ws_c.cell(i, 1, cr["variable_1"])
            ws_c.cell(i, 2, cr["variable_2"])
            ws_c.cell(i, 3, cr.get("methode", "Pearson"))
            ws_c.cell(i, 4, cr["r"])
            ws_c.cell(i, 5, cr["p_valeur"])
            ws_c.cell(i, 6, cr["interpretation"])
        ws_c.column_dimensions["A"].width = 26
        ws_c.column_dimensions["B"].width = 26
        ws_c.column_dimensions["F"].width = 40

    av = r.get("analyses_avancees") or {}

    if av.get("afe"):
        a = av["afe"]
        ws_afe = wb.create_sheet(l["aq_afe_titre"][3:][:31])
        ws_afe.cell(1, 1, l["aq_afe_kmo"]); ws_afe.cell(1, 2, f"{a['kmo_total']} ({a['kmo_interpretation']})")
        ws_afe.cell(2, 1, l["aq_afe_bartlett"]); ws_afe.cell(2, 2, f"χ²={a['bartlett_chi2']}, p={a['bartlett_p']}")
        ws_afe.cell(3, 1, l["aq_afe_facteurs"]); ws_afe.cell(3, 2, a["n_facteurs_extraits"])
        ws_afe.cell(4, 1, l["aq_afe_variance"]); ws_afe.cell(4, 2, f"{a['variance_expliquee_cumulee_pct']}%")
        for i in range(1, 5):
            ws_afe.cell(i, 1).font = Font(bold=True)
        ws_afe.cell(6, 1, l["aq_afe_charges"]).font = Font(bold=True)
        n_fact = a["n_facteurs_extraits"]
        _entete(ws_afe, ["Item"] + [f"Facteur {i+1}" for i in range(n_fact)], ligne=7)
        for i, c in enumerate(a["charges_factorielles"], start=8):
            ws_afe.cell(i, 1, c["item"])
            for fi, charge in enumerate(c["charges"]):
                ws_afe.cell(i, fi + 2, charge)
        ws_afe.column_dimensions["A"].width = 22

    if av.get("afc") and not av["afc"].get("erreur"):
        c = av["afc"]
        ws_afc = wb.create_sheet(l["aq_afc_titre"][3:][:31])
        for i, (label, val) in enumerate([("CFI", c["cfi"]), ("TLI", c["tli"]), ("RMSEA", c["rmsea"]), ("SRMR", c["srmr"])], start=1):
            ws_afc.cell(i, 1, label).font = Font(bold=True)
            ws_afc.cell(i, 2, val)
        ws_afc.cell(6, 1, c["interpretation"])
        _entete(ws_afc, ["Construit", "Item", "Charge", "p"], ligne=8)
        for i, ch in enumerate(c.get("charges_factorielles", []), start=9):
            ws_afc.cell(i, 1, ch["construit"]); ws_afc.cell(i, 2, ch["item"])
            ws_afc.cell(i, 3, ch["charge"]); ws_afc.cell(i, 4, ch["p_valeur"])
        ws_afc.column_dimensions["A"].width = 22

    if av.get("regressions"):
        ws_reg = wb.create_sheet(l["aq_reg_titre"][3:][:31])
        ligne = 1
        for reg in av["regressions"]:
            if reg.get("erreur"):
                continue
            ws_reg.cell(ligne, 1, f"{reg['dependante']} — R²={reg['r2']}, F={reg['f_stat']}, p={reg['f_p']}").font = Font(bold=True)
            ligne += 1
            _entete(ws_reg, [l["aq_reg_predicteur"], l["aq_reg_beta"], l["aq_reg_p"], "IC95% bas", "IC95% haut"], ligne=ligne)
            ligne += 1
            for p in reg["predicteurs"]:
                ws_reg.cell(ligne, 1, p["nom"] + (" *" if p["significatif"] else ""))
                ws_reg.cell(ligne, 2, p["beta"]); ws_reg.cell(ligne, 3, p["p_valeur"])
                ws_reg.cell(ligne, 4, p["ic95_bas"]); ws_reg.cell(ligne, 5, p["ic95_haut"])
                ligne += 1
            ligne += 1
        ws_reg.column_dimensions["A"].width = 24

    if av.get("mediations"):
        ws_med = wb.create_sheet(l["aq_med_titre"][3:][:31])
        _entete(ws_med, [l["aq_med_chemin"], l["aq_med_indirect"], "IC95% bas", "IC95% haut", l["aq_med_verdict"]])
        for i, m_ in enumerate([x for x in av["mediations"] if not x.get("erreur")], start=2):
            ws_med.cell(i, 1, f"{m_['independante']} → {m_['mediatrice']} → {m_['dependante']}")
            ws_med.cell(i, 2, m_["effet_indirect_a_x_b"])
            ws_med.cell(i, 3, m_["ic95_bas"]); ws_med.cell(i, 4, m_["ic95_haut"])
            ws_med.cell(i, 5, m_["type_mediation"] if m_["mediation_significative"] else "non significative")
        ws_med.column_dimensions["A"].width = 35

    synthese = analyse_data.get("synthese_interpretative")
    if synthese:
        ws_s = wb.create_sheet(l["aq_synthese_titre"][3:])
        ws_s.cell(1, 1, l["aq_synthese_avertissement"]).font = Font(italic=True, size=9, color="8A8574")
        ws_s.merge_cells(start_row=1, start_column=1, end_row=1, end_column=2)
        ws_s.row_dimensions[1].height = 40
        ligne = 3
        for titre_cle, contenu_cle in [
            ("aq_synthese_generale", "synthese_generale"), ("aq_synthese_fiabilite", "fiabilite_discussion"),
            ("aq_synthese_limites", "limites"), ("aq_synthese_recommandations", "recommandations"),
        ]:
            if not synthese.get(contenu_cle):
                continue
            ws_s.cell(ligne, 1, l[titre_cle]).font = Font(bold=True)
            ligne += 1
            ws_s.cell(ligne, 1, synthese[contenu_cle])
            ws_s.cell(ligne, 1).alignment = Alignment(wrap_text=True, vertical="top")
            ws_s.row_dimensions[ligne].height = 90
            ligne += 2
        ws_s.column_dimensions["A"].width = 100

        if synthese.get("tests_hypotheses"):
            ws_h = wb.create_sheet(l["aq_synthese_hypotheses"][:31])
            _entete(ws_h, [l["aq_hyp_code"], l["aq_hyp_verdict"], l["aq_hyp_justif"]])
            for i, th in enumerate(synthese["tests_hypotheses"], start=2):
                ws_h.cell(i, 1, th.get("code", ""))
                ws_h.cell(i, 2, th.get("verdict", ""))
                ws_h.cell(i, 3, th.get("justification", ""))
                ws_h.cell(i, 3).alignment = Alignment(wrap_text=True)
            ws_h.column_dimensions["B"].width = 30
            ws_h.column_dimensions["C"].width = 70
    elif analyse_data.get("synthese_statut") == "non_disponible":
        ws_s = wb.create_sheet(l["aq_synthese_titre"][3:])
        ws_s.cell(1, 1, l["aq_synthese_indisponible"])
        ws_s.column_dimensions["A"].width = 100

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf
